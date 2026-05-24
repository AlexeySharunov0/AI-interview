#!/usr/bin/env python3
"""Edit SharunovDiploma.docx per user requirements and methodology."""

from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = "SharunovDiploma.docx"
OUT = "SharunovDiploma.docx"

TOPIC = (
    "Разработка интеллектуального модуля оценки компетенций и анализа кода "
    "в реальном времени для платформы автоматизированного технического собеседования "
    "IT-специалистов"
)


def insert_after(paragraph, text="", style=None, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    para = Paragraph(new_p, paragraph._parent)
    if text:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = bold
    if style:
        para.style = style
    return para


def insert_many_after(paragraph, texts, style=None):
    last = paragraph
    for t in texts:
        last = insert_after(last, t, style=style)
    return last


def insert_code_after(paragraph, code_text):
    para = insert_after(paragraph, "")
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(12)
    pf = para.paragraph_format
    pf.left_indent = Cm(1.25)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    return para


def find_para(doc, startswith=None, equals=None, contains=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if equals and t == equals:
            return p
        if startswith and t.startswith(startswith):
            return p
        if contains and contains in t:
            return p
    return None


def replace_text_in_para(p, old, new):
    if old in p.text:
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
        if old in p.text and new not in p.text:
            p.text = p.text.replace(old, new)


def apply_formatting(doc):
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for level in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            h = doc.styles[level]
            h.font.name = "Times New Roman"
            h.font.size = Pt(14)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.first_line_indent = Cm(1.25)
            h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except KeyError:
            pass

    for p in doc.paragraphs:
        t = p.text.strip()
        if t in (
            "ВВЕДЕНИЕ",
            "ЗАКЛЮЧЕНИЕ",
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
            "ПРИЛОЖЕНИЕ А",
            "СОДЕРЖАНИЕ",
        ) or t.startswith(("1 АНАЛИЗ", "2 ПРОЕКТИРОВАНИЕ", "3 ТЕХНИКО")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
                r.bold = True
        elif t.startswith("Рисунок") or t.startswith("Таблица"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
                r.bold = False


def main():
    doc = Document(SRC)

    # Title topic
    p = find_para(doc, contains="на тему:")
    if p:
        replace_text_in_para(p, "на тему: « »", f"на тему: «{TOPIC}»")

    # Chapter renames
    p = find_para(doc, equals="1 АНАЛИТИЧЕСКАЯ ЧАСТЬ")
    if p:
        p.text = "1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ"

    p = find_para(doc, equals="2 ПРОЕКТНАЯ ЧАСТЬ")
    if p:
        p.text = "2 ПРОЕКТИРОВАНИЕ, РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ АВТОМАТИЗИРОВАННОЙ СИСТЕМЫ"

    # Renumber ch.2 tables (avoid duplicate 1,2 with chapter 1)
    renames = {
        "Таблица 1 - Сравнение платформ серверной разработки": "Таблица 3 - Сравнение платформ серверной разработки",
        "Таблица 2 - Сравнение решений для пользовательского интерфейса": "Таблица 4 - Сравнение решений для пользовательского интерфейса",
        "Таблица 3 - Сравнение СУБД для модуля собеседований": "Таблица 5 - Сравнение СУБД для модуля собеседований",
        "Таблица 4 - Сравнение инструментов разработки и лицензирования": "Таблица 6 - Сравнение инструментов разработки и лицензирования",
        "Таблица 5 - Функциональные требования к UI интеллектуального модуля": "Таблица 7 - Функциональные требования к UI интеллектуального модуля",
        "Таблица 6 - Основные сущности базы данных интеллектуального модуля": "Таблица 8 - Основные сущности базы данных интеллектуального модуля",
        "Таблица 7 - Сравнение вариантов реализации аутентификации": "Таблица 9 - Сравнение вариантов реализации аутентификации",
        "Таблица 8 - Результаты тестирования ключевых компонентов": "Таблица 10 - Результаты тестирования ключевых компонентов",
        "Рисунок 9 - ER-диаграмма базы данных модуля": "Рисунок 9 - Физическая ER-диаграмма базы данных модуля",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in renames:
            p.text = renames[t]

    # --- Bridge texts (insert AFTER anchor, process in reverse document order) ---
    bridges = [
        (
            "Таблица 2 - Сравнение платформ оценки IT-компетенций по ключевым критериям",
            [
                "Сопоставление показало, что ни одна из рассмотренных платформ не закрывает полный набор требований к разрабатываемой автоматизированной системе: replay-процесс, встроенный AI-интервьюер и комплексный античит должны быть реализованы в рамках единого контура. Полученные выводы использованы при формировании функциональных требований к интеллектуальному модулю и определении направления проектирования.",
            ],
        ),
        (
            "Рисунок 1 – схема выбранного технологического стека",
            [
                "На рисунке 1 представлена согласованная схема технологического стека: серверная часть (FastAPI, SQLAlchemy), клиентская (React, Vite), десктоп-оболочка (Electron), СУБД (PostgreSQL в продуктивном контуре, SQLite при локальной разработке), а также внешний LLM API для диалогового модуля. Такое сочетание обеспечивает единый сценарий работы кандидата и HR без дублирования бизнес-логики.",
            ],
        ),
        (
            "Таблица 6 - Сравнение инструментов разработки и лицензирования",
            [
                "Выбор бесплатных инструментов разработки и open-source компонентов снижает стоимость внедрения на этапе пилотной эксплуатации и соответствует ограничениям технико-экономического обоснования, представленного в главе 3.",
            ],
        ),
        (
            "Рисунок 2 - Архитектура взаимодействия модулей backend",
            [
                "Схема на рисунке 2 отражает потоки данных между модулями backend: оркестратор интервью координирует выдачу задач, приём решений, вызов code executor, запись метрик и античит-событий, а также обращения к AI-интервьюеру. Такое разделение позволяет независимо развивать подсистемы без переработки всей платформы.",
            ],
        ),
        (
            "Рисунок 3 - Структура каталогов backend-проекта",
            [
                "Структура каталогов backend (рисунок 3) организована по предметным модулям: auth, interviews, tasks, code_executor, ai_interviewer, metrics, anticheat, hr, reports. Каждый модуль содержит router и сервисную логику, что упрощает сопровождение и тестирование.",
            ],
        ),
        (
            "Рисунок 4 - Структура каталогов frontend-проекта",
            [
                "Клиентская часть (рисунок 4) построена на страничной модели React: отдельные экраны для кандидата (InterviewWorkspace, Results), HR (HRDashboard, HRReplay) и администрирования. Общение с API централизовано в модуле api/client.js.",
            ],
        ),
        (
            "Рисунок 5 - Структура и запуск desktop-части Electron",
            [
                "Desktop-контур (рисунок 5) использует Electron для загрузки собранного SPA из каталога electron/web, что обеспечивает кроссплатформенный запуск без изменения прикладного кода интерфейса.",
            ],
        ),
        (
            "Рисунок 6 - Интерфейс страницы Interview Workspace",
            [
                "Экран рабочего места кандидата (рисунок 6) объединяет условие задачи, редактор кода, панель тестов и чат с AI-интервьюером, минимизируя переключение контекста в ходе сессии.",
            ],
        ),
        (
            "Рисунок 7 - Интерфейс страницы HR Dashboard",
            [
                "Панель HR (рисунок 7) предоставляет сводную аналитику по кандидатам, фильтрацию по направлению и уровню, принятие решений hire/reject и переход к режиму воспроизведения сессии.",
            ],
        ),
        (
            "Рисунок 8 - Генерация Excel отчетов для HR",
            [
                "Модуль отчётности (рисунок 8) формирует документы Excel и PDF на основе данных интервью, что ускоряет передачу результатов в кадровые подразделения.",
            ],
        ),
        (
            "Рисунок 10 - Схема инфраструктуры Docker Compose",
            [
                "Инфраструктурная схема (рисунок 10) демонстрирует связь контейнеров PostgreSQL, Redis и backend-сервиса; healthcheck и зависимости сервисов заданы в docker-compose.yml.",
            ],
        ),
        (
            "Рисунок 11 - Фрагмент backend Dockerfile",
            [
                "Dockerfile backend (рисунок 11) фиксирует версию Python 3.11, установку зависимостей и команду запуска uvicorn, обеспечивая воспроизводимость сборки.",
            ],
        ),
        (
            "Рисунок 12 - Фрагмент docker-compose.yml",
            [
                "Фрагмент docker-compose.yml (рисунок 12) задаёт переменные окружения, тома данных и сетевое взаимодействие контейнеров.",
            ],
        ),
        (
            "Рисунок 13 - Сценарий прохождения интервью кандидатом",
            [
                "Блок-схема сценария (рисунок 13) отражает последовательность действий: авторизация, создание сессии, решение серии задач, диалог с AI, фиксация метрик и завершение с формированием отчёта для HR.",
            ],
        ),
        (
            "Рисунок 14 - Скрипт/листинг оркестратора интервью",
            [
                "Листинг оркестратора (рисунок 14) иллюстрирует ключевые методы create_interview, start_interview и submit_solution; полный текст приведён в приложении А.",
            ],
        ),
        (
            "Рисунок 15 - Скрипт/листинг code executor",
            [
                "Сервис выполнения кода (рисунок 15) реализует асинхронный запуск тестов с ограничением по времени; фрагмент run_tests приведён ниже и в приложении А.",
            ],
        ),
        (
            "Рисунок 16 - Скрипт/листинг AI interviewer service",
            [
                "AI-интервьюер (рисунок 16) формирует контекст задачи и истории переписки; при недоступности внешнего API возвращается резервное сообщение.",
            ],
        ),
        (
            "Рисунок 17 - Скрипт/листинг anticheat service",
            [
                "Сервис античита (рисунок 17) назначает веса событиям и вычисляет интегральный suspicion score для HR-аналитики.",
            ],
        ),
        (
            "Рисунок 18 - Схема авторизации JWT в модуле",
            [
                "Схема JWT (рисунок 18) показывает выдачу токена при входе, передачу в заголовке Authorization и проверку роли на защищённых маршрутах API.",
            ],
        ),
        (
            "Таблица 10 - Результаты тестирования ключевых компонентов",
            [
                "Результаты, сведённые в таблице 10, подтверждают работоспособность основного пользовательского сценария и согласованность модулей при совместной эксплуатации.",
            ],
        ),
        (
            "Рисунок 19 - Пример успешного прохождения тестов решения",
            [
                "На рисунке 19 показан результат прогона видимых тестов в рабочем месте кандидата; данный этап соответствует шагу 6 детального примера в п. 2.9.1.",
            ],
        ),
        (
            "Рисунок 20 - Пример отображения античит-событий в HR",
            [
                "Рисунок 20 демонстрирует отображение античит-событий и показателя подозрительности в интерфейсе HR при разборе сессии (шаг 9 примера в п. 2.9.1).",
            ],
        ),
    ]

    # Process bridges bottom-up by searching from end
    for anchor, texts in reversed(bridges):
        p = find_para(doc, equals=anchor)
        if p:
            insert_many_after(p, texts)

    # After table 1 ch1
    p = find_para(doc, equals="Таблица 1 - Сравнение подходов к оценке компетенций при техническом собеседовании")
    if p:
        # find last row - insert after table: get next non-empty para after table
        pass  # table is in doc.tables[0] - insert after paragraph before table content ends
    # Insert before table 1 - after intro text
    p = find_para(doc, contains="анализ кода в реальном времени")
    if p and "Таблица 1" not in p.text:
        insert_after(
            p,
            "Для систематизации сравнения подходов к оценке компетенций сведены ключевые характеристики в таблице 1. Данные использованы при обосновании выбора комплексного подхода с применением ИИ и трассировки процесса решения.",
        )

    # Between tables 3 and 4 in section 2.1.3
    p = find_para(doc, equals="Таблица 5 - Сравнение СУБД для модуля собеседований")
    if p:
        insert_after(
            p,
            "Сравнение СУБД (таблица 5) показало предпочтительность PostgreSQL для продуктивного развёртывания за счёт надёжной работы с JSON и транзакциями; на этапе разработки допускается SQLite для ускорения отладки. Далее в таблице 6 приведено сопоставление средств разработки.",
        )

    # Expand introduction tasks
    p = find_para(doc, contains="5. Функциональное тестирование")
    if p:
        insert_many_after(
            p,
            [
                "6. Проектирование базы данных на концептуальном, логическом и физическом уровнях.",
                "7. Проведение технико-экономического обоснования внедрения модуля.",
            ],
        )

    # === Section 2.5 DB design expansion ===
    p = find_para(doc, equals="2.5 Проектирование базы данных")
    if p:
        insert_many_after(
            p,
            [
                "Проектирование базы данных выполнялось в три этапа в соответствии с методологией проектирования информационных систем: концептуальный, логический и физический уровни. Такой подход обеспечивает преемственность от предметной области к реализации в СУБД и ORM SQLAlchemy.",
            ],
        )

    p = find_para(doc, contains="для повторной экспертной проверки спорных результатов")
    if p:
        insert_many_after(
            p,
            [
                "2.5.1 Концептуальное проектирование базы данных",
                "На концептуальном уровне предметная область описана сущностями: Пользователь, Компания, Вакансия, Интервью, Задача, Решение (отправка), Сообщение чата, Метрика, Событие античита, Уведомление, Тикет поддержки, Запись аудита. Связи: Компания — Пользователи (1:N); Пользователь — Интервью (1:N); Интервью — Задачи, Метрики, Античит-события, Сообщения (1:N); Задача — Отправки решений (1:N). Концептуальная модель отражает полный жизненный цикл сессии собеседования.",
                "[ВСТАВИТЬ РИСУНОК после данного абзаца: концептуальная диаграмма предметной области. Подпись: «Рисунок 9а — Концептуальная диаграмма предметной области».]",
                "2.5.2 Логическое проектирование базы данных",
                "На логическом уровне сущности отображены в таблицы реляционной модели с первичными и внешними ключами, кардинальностями и ограничениями целостности. Используются перечисления UserRole, InterviewStatus, DifficultyLevel, TaskDomain. Таблицы interviews, tasks, submissions обеспечивают трассировку оценки; metrics и anticheat_events — событийную аналитику.",
                "[ВСТАВИТЬ РИСУНОК: логическая ER-диаграмма с атрибутами и кардинальностями. Подпись: «Рисунок 9б — Логическая ER-диаграмма базы данных».]",
                "2.5.3 Физическое проектирование базы данных",
                "Физический уровень соответствует реализации в models.py и init_db.sql: типы столбцов, индексы по внешним ключам (interviews.user_id, tasks.interview_id, submissions.task_id), JSON-поля для наборов тестов и деталей событий. В продуктивном контуре применяется PostgreSQL; при локальной разработке — SQLite с тем же ORM-слоем.",
            ],
        )

    p = find_para(doc, contains="JSON-поля для наборов тестов")
    if p:
        insert_after(
            p,
            "Перед физической диаграммой (рисунок 9) приведена сводная таблица сущностей (таблица 8). Физическая ER-диаграмма отображает таблицы, типы данных, PK/FK и индексы в соответствии с models.py; полные листинги — в приложении А.",
        )

    # Code snippets in section 2.7
    p = find_para(doc, equals="2.7.1 Оркестрация сессии и адаптация сложности")
    if p:
        cur = insert_many_after(
            p,
            [
                "Класс InterviewOrchestrator управляет пятью задачами на сессию и адаптирует уровень сложности по результатам финальной отправки. Фрагмент реализации:",
            ],
        )
        code1 = """class InterviewOrchestrator:
    TASKS_PER_INTERVIEW = 5

    async def create_interview(self, db, user_id, level, specialty='developer'):
        interview = Interview(
            user_id=user_id, status=InterviewStatus.CREATED,
            level=level, specialty=specialty, total_tasks=TASKS_PER_INTERVIEW,
        )
        db.add(interview)
        await db.flush()
        return interview"""
        insert_code_after(cur, code1)

    p = find_para(doc, equals="2.7.2 Проверка кода и оценка")
    if p:
        cur = insert_many_after(
            p,
            [
                "Сервис CodeExecutor выполняет тесты в изолированном subprocess с таймаутом 10 с для Python и JavaScript:",
            ],
        )
        code2 = """async def run_tests(self, code, language, test_cases):
    if language not in self.SUPPORTED_LANGUAGES:
        return {"passed": 0, "total": len(test_cases), "error": "Язык не поддерживается"}
    results = []
    for tc in test_cases:
        result = await self._run_single_test(code, language, tc["input"], tc["expected"])
        results.append(result)
    return {"passed": sum(r["passed"] for r in results), "total": len(test_cases), "results": results}"""
        insert_code_after(cur, code2)

    p = find_para(doc, equals="2.7.3 Диалоговый AI-контур")
    if p:
        cur = insert_many_after(
            p,
            [
                "При отправке сообщения кандидата AI-интервьюер сохраняет реплику, собирает контекст задачи и вызывает LLM; при ошибке API возвращается резервный ответ:",
            ],
        )
        code3 = """ai_response = await scibox_client.chat(messages=messages, system_prompt=SYSTEM_PROMPT + task_ctx)
if not ai_response:
    ai_response = "AI-интервьюер временно недоступен. Попробуйте позже."
db.add(ChatMessage(interview_id=interview_id, task_id=task_id, sender="ai", content=ai_response))"""
        insert_code_after(cur, code3)

    p = find_para(doc, equals="2.7.4 Античит и replay-аналитика")
    if p:
        cur = insert_many_after(
            p,
            [
                "Расчёт интегрального показателя подозрительности выполняется по сумме весов событий:",
            ],
        )
        code4 = """SEVERITY = {"paste": 0.5, "tab_switch": 0.2, "devtools": 0.7, "rapid_code": 0.8}
base = min(total_severity / 3.0, 0.8)
repeat_penalty = max(0, (event_count - 2) * 0.05)
return min(base + repeat_penalty, 1.0)"""
        insert_code_after(cur, code4)

    # Detailed example 2.9.1 — insert in direct order before «2.10 Выводы»
    p = find_para(doc, equals="2.10 Выводы по проектной части")
    if p:
        example = [
            "2.9.1 Детальный пример прохождения интервью",
            "Рассмотрим сквозной сценарий для кандидата с ролью candidate (учётная запись candidate3@gmail.com), направление «Разработка», уровень middle.",
            "Шаг 1. Авторизация. Кандидат вводит email и пароль на странице Login; frontend отправляет POST /api/auth/login, получает JWT и сохраняет токен в localStorage. Дальнейшие запросы содержат заголовок Authorization: Bearer <token>.",
            "Шаг 2. Создание сессии. На InterviewSelector выбираются specialty=developer и level=middle; вызывается POST /api/interviews/. InterviewOrchestrator создаёт запись interviews со статусом created и total_tasks=5.",
            "Шаг 3. Старт первой задачи. POST /api/interviews/{id}/start переводит статус в task_active; task_generator выбирает шаблон из task_templates (например, задача на массивы), при необходимости перефразирует описание через LLM и создаёт запись tasks с visible_tests и hidden_tests.",
            "Шаг 4. Работа в InterviewWorkspace. Каждые 30 с frontend отправляет POST /api/metrics/event с типом code_snapshot. При вставке из буфера фиксируется POST /api/anticheat/event с типом paste.",
            "Шаг 5. Прогон тестов. Кандидат нажимает «Запустить тесты»; отправляется решение без is_final; code_executor прогоняет visible_tests, возвращает passed/total и детализацию по каждому тесту.",
            "Шаг 6. Диалог с AI. В чате POST /api/chat/send; AIInterviewer добавляет сообщения в chat_messages, формирует контекст задачи и возвращает наводящий ответ без готового кода (см. рисунок 19).",
            "Шаг 7. Финальная отправка. submit с is_final=true запускает видимые и скрытые тесты; score сохраняется в submissions; orchestrator увеличивает completed_tasks, адаптирует level и выдаёт следующую задачу или завершает интервью.",
            "Шаг 8. Завершение. После пятой задачи статус interviews → completed, рассчитывается total_score; кандидат видит Results, HR получает запись в дашборде.",
            "Шаг 9. Разбор HR. HR открывает HRReplay: timeline объединяет snapshots, chat_messages, anticheat_events и submissions; при необходимости экспортируется PDF/Excel (рисунок 20).",
            "Таким образом, пример демонстрирует согласованную работу всех подсистем автоматизированной системы в одном пользовательском сценарии.",
        ]
        parent = p._p.getparent()
        anchor_idx = parent.index(p._p)
        for offset, text in enumerate(example):
            new_p = OxmlElement("w:p")
            parent.insert(anchor_idx + offset, new_p)
            para = Paragraph(new_p, p._parent)
            run = para.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            if text.startswith("2.9.1"):
                run.bold = True

    # Chapter 3, Conclusion, Sources, Appendix - after last paragraph
    last = doc.paragraphs[-1]
    tail = [
        ("3 ТЕХНИКО-ЭКОНОМИЧЕСКОЕ ОБОСНОВАНИЕ", True),
        (
            "Технико-экономическое обоснование подтверждает целесообразность разработки и пилотного внедрения автоматизированной системы технического собеседования с интеллектуальным модулем оценки компетенций.",
            False,
        ),
        ("3.1 Обоснование выбора программно-технологического стека", True),
        (
            "Выбранный стек FastAPI, React, Electron, PostgreSQL/SQLite и open-source инструменты разработки минимизирует лицензионные затраты и обеспечивает масштабирование без смены архитектуры. По сравнению с SaaS-платформами (HackerRank, Codility) собственное решение даёт контроль над хранением персональных данных и развитием античита и replay.",
            False,
        ),
        ("3.2 Оценка затрат на разработку и внедрение", True),
        (
            "В структуру затрат включены трудозатраты разработчика (проектирование, реализация, тестирование — ориентировочно 4–6 человеко-месяцев на этапе ВКР), инфраструктура (VPS/облако 1500–3000 руб./мес., домен, SSL), переменные расходы на LLM API (по факту запросов в чате). На этапе учебной апробации возможно использование бесплатных тарифов и локального развёртывания.",
            False,
        ),
        ("Таблица 11 - Оценка затрат на разработку и эксплуатацию", False),
        (
            "В таблице 11 приведены ориентировочные статьи затрат; итоговая сумма на год пилотной эксплуатации составляет порядка 120–180 тыс. руб. при частичной удалённой инфраструктуре и минимальном объёме платных AI-запросов. Методика расчёта: суммирование единовременных затрат на разработку (труд) и помесячных эксплуатационных расходов за 12 месяцев.",
            False,
        ),
        ("3.3 Оценка экономического эффекта от внедрения", True),
        (
            "Экономический эффект выражается в сокращении времени senior-разработчиков на первичные интервью (оценочно 2–4 часа на вакансию), ускорении отбора и снижении доли повторных собеседований за счёт объективной трассировки процесса. Качественный эффект — стандартизация критериев оценки и прозрачность решений для HR.",
            False,
        ),
        ("3.4 Выводы по технико-экономическому обоснованию", True),
        (
            "Разработка автоматизированной системы экономически целесообразна для организаций с регулярным потоком IT-вакансий; при росте числа интервью удельные затраты на одну сессию снижаются.",
            False,
        ),
        ("ЗАКЛЮЧЕНИЕ", True),
        (
            "В дипломном проекте разработан интеллектуальный модуль оценки компетенций и анализа кода в реальном времени в составе платформы автоматизированного технического собеседования IT-специалистов.",
            False,
        ),
        (
            "Выполнен анализ предметной области и существующих решений, сформированы функциональные и нефункциональные требования, обоснован технологический стек, спроектирована архитектура и база данных на трёх уровнях, реализованы ключевые сервисы (оркестратор, выполнение кода, AI-интервьюер, метрики, античит, HR-инструменты), проведено функциональное тестирование и технико-экономическое обоснование.",
            False,
        ),
        (
            "Направления развития: полноценная sandbox-изоляция кода в Docker, интеграция VPN-модуля, подключение PostgreSQL в продуктивном контуре, расширение банка задач и языков программирования.",
            False,
        ),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", True),
        ("1. ГОСТ 7.32–2017. Отчёт о научно-исследовательской работе. Структура и правила оформления.", False),
        ("2. FastAPI Documentation [Электронный ресурс]. – URL: https://fastapi.tiangolo.com (дата обращения: 01.05.2026).", False),
        ("3. React Documentation [Электронный ресурс]. – URL: https://react.dev (дата обращения: 01.05.2026).", False),
        ("4. SQLAlchemy Documentation [Электронный ресурс]. – URL: https://docs.sqlalchemy.org (дата обращения: 01.05.2026).", False),
        ("5. PostgreSQL Documentation [Электронный ресурс]. – URL: https://www.postgresql.org/docs/ (дата обращения: 01.05.2026).", False),
        ("6. Electron Documentation [Электронный ресурс]. – URL: https://www.electronjs.org/docs (дата обращения: 01.05.2026).", False),
        ("7. OWASP API Security Top 10 [Электронный ресурс]. – URL: https://owasp.org (дата обращения: 01.05.2026).", False),
        ("8. HackerRank. Technical Hiring Platform [Электронный ресурс]. – URL: https://www.hackerrank.com (дата обращения: 01.05.2026).", False),
        ("9. Codility. Technical Assessment Platform [Электронный ресурс]. – URL: https://www.codility.com (дата обращения: 01.05.2026).", False),
        ("10. Fowler M. Patterns of Enterprise Application Architecture. – Addison-Wesley, 2002.", False),
        ("ПРИЛОЖЕНИЕ А", True),
        (
            "Листинги программного кода автоматизированной системы (фрагменты модулей backend/app/interviews/orchestrator.py, code_executor/executor.py, ai_interviewer/service.py, anticheat/service.py, models/models.py).",
            False,
        ),
        (
            "[ВСТАВИТЬ ПОЛНЫЕ ЛИСТИНГИ или использовать рисунки 14–17 главы 2; шрифт Courier New, 12 пт, формат А4.]",
            False,
        ),
    ]
    cur = last
    for text, is_heading in tail:
        cur = insert_after(cur, text)
        if is_heading:
            cur.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cur.runs:
                r.bold = True

    # Add TOC before introduction
    p_intro = find_para(doc, equals="ВВЕДЕНИЕ")
    if p_intro:
        toc_lines = [
            "СОДЕРЖАНИЕ",
            "ВВЕДЕНИЕ ................................................................................................ 4",
            "1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ ........................................................ 6",
            "2 ПРОЕКТИРОВАНИЕ, РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ АВТОМАТИЗИРОВАННОЙ СИСТЕМЫ ........ 12",
            "3 ТЕХНИКО-ЭКОНОМИЧЕСКОЕ ОБОСНОВАНИЕ .................................................. 38",
            "ЗАКЛЮЧЕНИЕ ........................................................................................... 42",
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ ................................................... 43",
            "ПРИЛОЖЕНИЕ А ....................................................................................... 44",
        ]
        parent = p_intro._p.getparent()
        for line in reversed(toc_lines):
            new_p = OxmlElement("w:p")
            parent.insert(parent.index(p_intro._p), new_p)
            para = Paragraph(new_p, p_intro._parent)
            run = para.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            if line == "СОДЕРЖАНИЕ":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
            else:
                para.paragraph_format.first_line_indent = Cm(0)

    # Table 11 - data as text rows after caption (user may convert to Word table)
    p = find_para(doc, equals="Таблица 11 - Оценка затрат на разработку и эксплуатацию")
    if p:
        insert_many_after(
            p,
            [
                "Разработка (труд) | 80 000–120 000 | —",
                "Сервер/облако | 5 000 | 18 000–36 000",
                "LLM API | — | 6 000–24 000",
                "Прочее (домен, SSL) | 2 000 | 3 000",
                "Итого | 87 000–127 000 | 27 000–63 000",
            ],
        )

    apply_formatting(doc)
    doc.save(OUT)
    print(f"Saved {OUT}, paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
