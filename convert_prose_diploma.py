#!/usr/bin/env python3
"""Convert bullet/numbered lists to prose; complete title page."""

import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = "/Users/alexeysharunov/Downloads/SharunovDiploma (1).docx"
OUT = "/Users/alexeysharunov/Downloads/v2 2/SharunovDiploma_исправленный.docx"

TOPIC = (
    "Разработка интеллектуального модуля оценки компетенций и анализа кода "
    "в реальном времени для платформы автоматизированного технического "
    "собеседования IT-специалистов"
)

TITLE_BLOCK = [
    "Выполнил студент группы 404ИС-22",
    "",
    "___________________",
    "/А.А. Шарунов/",
    "Руководитель",
    "",
    "___________________",
    "/Ю.Г. Грибова/",
    "Нормоконтролер",
    "",
    "___________________",
    "/Ю.Г. Грибова/",
    "Консультант по технико-экономическому обоснованию проекта",
    "",
    "___________________",
    "/С.В. Бармина/",
    "Старший консультант",
    "",
    "___________________",
    "/Ю.Г. Грибова/",
]

INTRO_TASKS_PROSE = (
    "Для достижения поставленной цели необходимо провести анализ существующих решений "
    "в области автоматизированной оценки IT-компетенций и анализа кода, сформировать "
    "и описать функциональные возможности интеллектуального модуля, обосновать выбор "
    "технологического стека и спроектировать архитектуру, реализовать сервис выполнения "
    "кода, сбор метрик, AI-интервьюер, генератор задач и античит, спроектировать базу "
    "данных на концептуальном, логическом и физическом уровнях, провести функциональное "
    "тестирование и технико-экономическое обоснование внедрения модуля."
)

SECTION_HEADING = re.compile(
    r"^(\d+(\.\d+)*\s+[А-ЯA-Z])|^(ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|СПИСОК|ПРИЛОЖЕНИЕ|СОДЕРЖАНИЕ|Таблица|Рисунок)|^где\s"
)


def is_list_item(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if t.startswith("- ") or t.startswith("– "):
        return True
    if t.startswith("Шаг "):
        return True
    if SECTION_HEADING.match(t):
        return False
    if re.match(r"^\d+\.\d+", t):
        return False
    if re.match(r"^\d+\.\s", t) and not re.match(r"^\d+\.\s*$", t):
        # numbered list like "1. Проведение..."
        if re.match(r"^\d+\.\s+[а-яё]", t, re.I):
            return True
        if re.match(r"^\d+\.\s+[А-ЯЁ]", t) and len(t) > 25:
            return True
    return False


def clean_item(text: str) -> str:
    t = text.strip()
    t = re.sub(r"^[-–]\s*", "", t)
    t = re.sub(r"^\d+\.\s*", "", t)
    t = re.sub(r"^Шаг\s+\d+\.\s*", "", t)
    if t and t[-1] not in ".;:":
        t += "."
    return t[0].lower() + t[1:] if t else t


def merge_list_items(items: list[str]) -> str:
    parts = [clean_item(x) for x in items if x.strip()]
    if not parts:
        return ""
    text = " ".join(parts)
    text = re.sub(
        r"\.\s+([а-яё])",
        lambda m: ". " + m.group(1).upper(),
        text,
    )
    return text[0].upper() + text[1:]


# Точечные замены после автослияния списков
POLISH_BY_PREFIX = [
    (
        "Наличие автоматизированной проверки кода кандидата",
        "При сравнении платформ учитываются наличие автоматизированной проверки кода "
        "(видимые и скрытые тесты, автоскоринг), поддержка оценки процесса решения "
        "(история и снимки кода, режим воспроизведения), диалоговое взаимодействие с "
        "AI-интервьюером, интеграция с HR-процессами и отчётностью, локализация на "
        "русском языке, безопасность исполнения и античит, а также масштабируемость "
        "и работа в реальном времени при нагрузке.",
    ),
    (
        "Функциональные требования.",
        "Функциональные требования к модулю включают выполнение кода кандидата в "
        "изолированной среде с ограничениями по времени и памяти (Python и JavaScript), "
        "проверку решений по видимым и скрытым тестам с детализированными результатами, "
        "периодический сбор снимков кода для анализа и воспроизведения сессии, диалог с "
        "AI-интервьюером (ответы на вопросы и наводящие подсказки без раскрытия решения), "
        "гибридную генерацию задач из банка шаблонов с возможной модификацией через ИИ, "
        "а также фиксацию подозрительных событий с расчётом показателя подозрительности.",
    ),
    (
        "Проверка решений по наборам видимых",
        "",  # удалить дубль после слияния
    ),
    (
        "Производительность: время отклика API",
        "Дополнительно к измеримым показателям п. 1.2.2 обеспечиваются производительность "
        "(время отклика API при выполнении кода), масштабируемость при множестве одновременных "
        "собеседований, надёжность за счёт резервного механизма при недоступности AI API и "
        "безопасность за счёт изолированного выполнения кода и валидации входных данных.",
    ),
    (
        "поддержка работы в реальном времени без деградации",
        "При выборе стека учитывались поддержка работы в реальном времени без деградации UX, "
        "безопасный запуск пользовательского кода с фиксацией результатов тестов, интеграция "
        "с AI-сервисом для диалога и генерации вариаций задач, инструменты replay для HR и "
        "минимальные лицензионные издержки на этапе пилотной эксплуатации.",
    ),
    (
        "создание интервью и корректная выдача первой задачи",
        "Функциональное тестирование охватывает создание интервью и выдачу первой задачи по "
        "домену и уровню, проверку решений на видимых и скрытых тестах с сохранением артефактов, "
        "работу AI-контура в штатном и fallback-режиме, формирование итоговой оценки в "
        "HR-интерфейсе и воспроизведение timeline с синхронизацией чата, кода и античит-событий.",
    ),
]


def insert_before(paragraph, text, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    para = Paragraph(new_p, paragraph._parent)
    if text:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = bold
    return para


def title_has_student_block(doc):
    for p in doc.paragraphs:
        if "Шарунов" in p.text or "404ИС" in p.text:
            return True
    return False


def fix_title_page(doc):
    for p in doc.paragraphs:
        if "на тему:" in p.text:
            p.text = f"на тему: «{TOPIC}»"
        if p.text.strip() == "на тему: « »":
            p.text = f"на тему: «{TOPIC}»"

    if title_has_student_block(doc):
        return

    moscow = None
    for p in doc.paragraphs:
        if p.text.strip() == "МОСКВА":
            moscow = p
            break
    if not moscow:
        return

    anchor = moscow
    for line in reversed(TITLE_BLOCK):
        anchor = insert_before(anchor, line)


def replace_intro_tasks(doc):
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "Для достижения цели сформулирован ряд задач:":
            start = i
        elif start is not None and end is None and t.startswith("Объектом исследования"):
            end = i
            break
    if start is None or end is None:
        return
    doc.paragraphs[start].text = INTRO_TASKS_PROSE
    for j in range(end - 1, start, -1):
        el = doc.paragraphs[j]._element
        el.getparent().remove(el)


def convert_list_groups(doc):
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not is_list_item(t):
            i += 1
            continue
        group = []
        j = i
        while j < len(doc.paragraphs) and is_list_item(doc.paragraphs[j].text):
            group.append(doc.paragraphs[j].text)
            j += 1
        if len(group) >= 2:
            merged = merge_list_items(group)
            doc.paragraphs[i].text = merged
            for k in range(j - 1, i, -1):
                el = doc.paragraphs[k]._element
                el.getparent().remove(el)
            i += 1
        else:
            # single dash item -> inline sentence
            if t.startswith("- ") or t.startswith("– "):
                doc.paragraphs[i].text = clean_item(t)
                doc.paragraphs[i].text = (
                    doc.paragraphs[i].text[0].upper() + doc.paragraphs[i].text[1:]
                )
            i += 1


def polish_paragraphs(doc):
    to_remove = []
    for p in doc.paragraphs:
        t = p.text.strip()
        for prefix, replacement in POLISH_BY_PREFIX:
            if not prefix or not t.startswith(prefix):
                continue
            if replacement == "":
                to_remove.append(p)
            else:
                p.text = replacement
            break
    for p in to_remove:
        el = p._element
        el.getparent().remove(el)


def remove_list_intro_lines(doc):
    intros = (
        "Критерии выбора программных решений в рамках подтемы:",
        "Ключевые архитектурные требования к интеллектуальному модулю:",
        "Практические преимущества выбранного backend-стека:",
        "Требования к пользовательскому интерфейсу при проектировании:",
        "Логика адаптации уровня в рамках реализованного подхода:",
        "События, влияющие на интегральный показатель подозрительности:",
        "Базовые меры безопасности, реализованные в проекте:",
        "Основные сценарии, покрытые функциональным тестированием:",
        "Итоги разработки интеллектуального модуля по подтеме:",
        "Для достижения цели сформулирован ряд задач:",
        "Функциональные требования.",
        "Функциональные требования интеллектуального модуля",
    )
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in intros or t.endswith(":") and t in intros:
            p.text = ""


def enhance_intro(doc):
    """Add substance to introduction if methods/structure missing."""
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() != "ВВЕДЕНИЕ":
            continue
        nxt = None
        if idx + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[idx + 1]
        if nxt and "метод" not in nxt.text.lower() and "структур" not in nxt.text.lower():
            extra = (
                "Работа выполнена на материалах разработанной платформы технического "
                "собеседования (backend на FastAPI, клиент на React, desktop-оболочка Electron). "
                "Методы исследования включают сравнительный анализ аналогов, проектирование "
                "архитектуры и базы данных, программную реализацию модулей, функциональное "
                "тестирование и технико-экономическое обоснование. Структура дипломного проекта: "
                "введение, анализ предметной области, проектирование и реализация системы, "
                "технико-экономическое обоснование, заключение, список источников и приложение."
            )
            new_p = OxmlElement("w:p")
            nxt._p.addprevious(new_p)
            para = Paragraph(new_p, nxt._parent)
            run = para.add_run(extra)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        break


def main():
    doc = Document(SRC)
    fix_title_page(doc)
    replace_intro_tasks(doc)
    convert_list_groups(doc)
    polish_paragraphs(doc)
    remove_list_intro_lines(doc)
    enhance_intro(doc)
    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
